from __future__ import annotations

import json
import shutil
import subprocess
import tempfile
from pathlib import Path
from uuid import uuid4
from typing import Any

import pdfplumber
import pytesseract
from fastapi import HTTPException, UploadFile
from loguru import logger
from PIL import Image
from slugify import slugify
from sqlalchemy.orm import Session

from app.core.config import settings
from app.models.doc_claim_community_forest_resource import DocClaimCommunityForestResource
from app.models.doc_claim_community_rights import DocClaimCommunityRights
from app.models.doc_claim_forest_land import DocClaimForestLand
from app.models.doc_title_community_forest_resources import DocTitleCommunityForestResources
from app.models.doc_title_community_forest_rights import DocTitleCommunityForestRights
from app.models.doc_title_under_occupation import DocTitleUnderOccupation
from app.models.master_document import MasterDocument
from app.schemas.document import DocumentMetadata, DocumentUploadResponse, RuleBasedPayload

TEMPLATE_MODEL_MAP = {
  "DOC_CLAIM_FOREST_LAND": DocClaimForestLand,
  "DOC_CLAIM_COMMUNITY_RIGHTS": DocClaimCommunityRights,
  "DOC_CLAIM_COMMUNITY_FOREST_RESOURCE": DocClaimCommunityForestResource,
  "DOC_TITLE_UNDER_OCCUPATION": DocTitleUnderOccupation,
  "DOC_TITLE_COMMUNITY_FOREST_RIGHTS": DocTitleCommunityForestRights,
  "DOC_TITLE_COMMUNITY_FOREST_RESOURCES": DocTitleCommunityForestResources,
}


class DocumentIngestionService:
  @staticmethod
  def _persist_upload(file: UploadFile) -> Path:
    uploads_dir = settings.uploads_path
    safe_name = slugify(Path(file.filename or "document").stem)
    dest_name = f"{safe_name}-{uuid4().hex}{Path(file.filename or '').suffix or '.bin'}"
    destination = uploads_dir / dest_name
    with destination.open("wb") as buffer:
      shutil.copyfileobj(file.file, buffer)
    return destination

  @staticmethod
  def _extract_text(file_path: Path) -> str:
    suffix = file_path.suffix.lower()
    if suffix in {".pdf"}:
      with pdfplumber.open(file_path) as pdf:
        pages = [page.extract_text() or "" for page in pdf.pages]
      text = "\n".join(pages).strip()
      if text:
        return text
      return DocumentIngestionService._run_ocr(file_path)
    if suffix in {".png", ".jpg", ".jpeg", ".tif"}:
      return DocumentIngestionService._run_ocr(file_path)
    if suffix in {".txt"}:
      return file_path.read_text(encoding="utf-8", errors="ignore")
    if suffix in {".doc", ".docx"}:
      try:
        import docx
      except ImportError as exc:
        raise HTTPException(status_code=500, detail="python-docx is required") from exc
      doc = docx.Document(str(file_path))
      return "\n".join(p.text for p in doc.paragraphs)
    return file_path.read_text(encoding="utf-8", errors="ignore")

  @staticmethod
  def _run_ocr(file_path: Path) -> str:
    image = Image.open(file_path)
    return pytesseract.image_to_string(image, lang=settings.ocr_language_hint)

  @staticmethod
  def _invoke_rule_based(text: str, enable_translation: bool) -> RuleBasedPayload:
    script = settings.rule_based_script
    if not script.exists():
      raise HTTPException(status_code=500, detail="Rule-based module not found")

    with tempfile.TemporaryDirectory() as tmpdir:
      input_path = Path(tmpdir) / "input.txt"
      output_path = Path(tmpdir) / "result.json"
      input_path.write_text(text, encoding="utf-8")

      cmd = [
        "python",
        str(script),
        "--input",
        str(input_path),
        "--output",
        str(output_path),
      ]
      if enable_translation:
        cmd.append("--translate")

      logger.info("Executing rule-based pipeline: {}", " ".join(cmd))
      proc = subprocess.run(cmd, cwd=settings.rule_based_recog_dir, check=False, capture_output=True, text=True)
      if proc.returncode != 0:
        logger.error("Rule-based pipeline failed: {}", proc.stderr)
        raise HTTPException(status_code=500, detail="Rule-based pipeline failed")

      raw_payload = json.loads(output_path.read_text(encoding="utf-8"))
      if isinstance(raw_payload, list):
        if not raw_payload:
          raise HTTPException(status_code=500, detail="Empty rule-based response")
        raw_payload = raw_payload[0]
      return RuleBasedPayload.model_validate(raw_payload)

  @staticmethod
  def _persist_entities(db: Session, template_id: str, document_id: int, entities: dict[str, Any]):
    model = TEMPLATE_MODEL_MAP.get(template_id)
    if not model:
      raise HTTPException(status_code=400, detail=f"Unsupported template: {template_id}")
    allowed_columns = {column.name for column in model.__table__.columns}
    payload = {k.lower(): v for k, v in entities.items() if k.lower() in allowed_columns}
    instance = model(document_id=document_id, **payload)
    db.add(instance)

  def ingest(self, db: Session, file: UploadFile, metadata: DocumentMetadata) -> DocumentUploadResponse:
    saved_path = self._persist_upload(file)
    raw_text = self._extract_text(saved_path)

    document = MasterDocument(
      file_name=file.filename or saved_path.name,
      file_path=str(saved_path),
      document_type=metadata.document_type,
      language=metadata.language,
      raw_text=raw_text,
      processing_status="PROCESSING",
    )

    db.add(document)
    db.commit()
    db.refresh(document)

    try:
      payload = self._invoke_rule_based(raw_text, settings.enable_translation)
      document.processing_status = "COMPLETED"
      document.document_type = payload.template_id
      document.extracted_payload = payload.model_dump_json()
      self._persist_entities(db, payload.template_id, document.id, payload.entities)
      db.add(document)
      db.commit()
    except Exception:
      db.rollback()
      document.processing_status = "FAILED"
      db.add(document)
      db.commit()
      raise

    return DocumentUploadResponse(
      document_id=document.id,
      document_type=document.document_type,
      processing_status=document.processing_status,
      template_id=document.document_type,
      created_at=document.upload_timestamp,
    )


document_ingestion_service = DocumentIngestionService()
